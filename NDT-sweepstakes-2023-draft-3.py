### import statements
import pandas as pd
pd.options.mode.chained_assignment = None  # supress warnings, i've tested the code
import numpy as np
from enum import Enum
import os
import re
import docx
from docxcompose.composer import Composer
import argparse

# take input for year and season
command_line_argument_parser=argparse.ArgumentParser()
command_line_argument_parser.add_argument("-y","--year",help="Year of report to generate, default 2023",type=int,default=2023)
command_line_argument_parser.add_argument("-s","--season",help="Season to generate, fall or spring, default fall",type=str,default='f',choices=['fall','spring','f','s'])
command_line_argument_parser.add_argument("-n","--no_report",help="Disable docx report generation",action='store_true')
command_line_argument_parser.add_argument("-d","--debug",help="Debug mode",action='store_true')
command_line_argument_parser.add_argument("-v","--validate",help="Validate tournaments and schools",action='store_true')
command_line_argument_parser.add_argument("-R","--separate_rr",help="Separate round-robins from varsity tournaments",action='store_true')
arguments=command_line_argument_parser.parse_args()

NO_REPORT_GEN=arguments.no_report
VALIDATION_MODE=arguments.validate
SEPARATE_ROUNDROBINS = arguments.separate_rr
YEAR_TO_PROCESS=arguments.year
REPORTS_GENERATED_FOLDER='generated_reports/'+str(YEAR_TO_PROCESS)+'/'
PREVIOUS_YEAR_REPORTS_FOLDER='generated_reports/'+str(YEAR_TO_PROCESS-1)+'/'
if (arguments.season=='f')|(arguments.season=='fall'):
    REPORT_TO_GENERATE = 1
else:
    REPORT_TO_GENERATE = 2



### global definitions
def ndt_points_from_prelims(prelim_percentage): ## taken from the ranking procedure.
    points = 8
    points += prelim_percentage>0
    points += prelim_percentage>=0.21
    points += prelim_percentage>=0.33
    points += prelim_percentage>0.4999
    points += prelim_percentage>0.50
    points += prelim_percentage>=0.67
    points += prelim_percentage>=0.80
    points += prelim_percentage>0.9999
    return points

def print_if_debug(string):
    if arguments.debug:
        print(string)
    return

# I might be able to significantly speed this code up by attempting to vectorize, but do I really care?
def ndt_winner_points_from_elims(loser_ballots): # taken from ranking procedure: unanimous elim wins (or byes) are worth 6, else 5
    return 6-(loser_ballots>=1)
def ndt_loser_points_from_elims(loser_ballots): # taken from ranking procedure: Showing up is worth 3 points, taking a ballot worth 4
    return 3+(loser_ballots>=1)

class Division(Enum):
    VARSITY = 'v'
    JUNIOR_VARSITY = 'jv'
    NOVICE = 'n'
    ROUND_ROBIN = 'rr'
    TOTAL = 'total' ## hack. TODO: Remove.

    
# I could probably get away without referencing the year, but I want to be able to process last year's results to generate
# this year's spring reports with 'movers' and 'new schools' so I plan for the future
class tournament():
        def __init__(self,tournament_name,tournament_year,prelim_count_vector,division_vector):
            self.prelim_counts = prelim_count_vector
            self.divisions = division_vector
            self.name = tournament_name
            self.year = tournament_year

MINIMUM_SCHOOLS_PER_DIVISION = 3 # taken from the ranking procedure. all of these are checked inclusively.
MINIMUM_TEAMS_PER_DIVISION = 6
MINIMUM_PRELIMS_PER_DIVISION = 4
MAXIMUM_ENTRIES_COUNTED_PER_SCHOOL = 2
MAXIMUM_TOURNAMENTS_COUNTED_PER_SCHOOL = 8
NEW_SCHOOL_POINTS_THRESHOLD = 32 # used in spring reports only
MOVERS_THRESHOLD = 32 # used in spring reports only


NDT_DISTRICTS = range(1,9,1) # 1-8, not inclusive. Unlikely to change, but i'll centralize it anyway

def first_or_second(): # clunky, but avoids hard-coding.
    if REPORT_TO_GENERATE==1:
        ordinal="first"
        season="fall"
        report_year_string=str(YEAR_TO_PROCESS)
    else:
        ordinal="second"
        season="spring"
        report_year_string=str(YEAR_TO_PROCESS+1)
    return [ordinal,season,report_year_string]


[report_ordinal,report_season,report_year] = first_or_second()

## Prepare to replace school names with 'pretty' school names for display: 'Minnesota' -> 'University of Minnesota'

school_alias_dataframe=pd.read_csv('school-alias-map.csv')
school_alias_dataframe['Alias-0']=school_alias_dataframe['Display-School'] # Some schools are already listed with their pretty name.
school_alias_dict_dataframe=pd.DataFrame()
for alias_item in school_alias_dataframe.columns[1:]:
    temp_school_alias_dataframe=pd.DataFrame()
    temp_school_alias_dataframe[['Display Name','Alias']]=school_alias_dataframe[['Display-School',alias_item]]
    temp_school_alias_dataframe.dropna(inplace=True)
    if school_alias_dict_dataframe.empty:
        school_alias_dict_dataframe = temp_school_alias_dataframe
    else:
        school_alias_dict_dataframe = pd.concat([school_alias_dict_dataframe,temp_school_alias_dataframe],ignore_index=True)

school_alias_dict=school_alias_dict_dataframe.set_index('Alias')['Display Name'].T.to_dict()

def apply_dictionary_to_results_dataframe(results_dataframe,school_dictionary):
    results_dataframe_test = results_dataframe['School'].map(school_dictionary)
    unmapped_schools=results_dataframe[results_dataframe_test.isna()]
    unmapped_school_count=len(unmapped_schools.index)
    if unmapped_school_count>0:
        print('The following rows contain unmapped schools:')
        print(unmapped_schools['School'].to_string())
        raise Exception('There are unmapped schools')
    results_dataframe['School'] = results_dataframe['School'].map(school_dictionary).fillna(results_dataframe['School'])
    return results_dataframe




# may properly drop invalid divisions, will certainly at least error out if presented with an invalid division.
# does not properly consider prelim seed in first elim round
# does not account for 'extenuating circumstances', II.(g), needs more code

# I use '3-0' to describe any unanimous win (bye, forfeit, walkover). If the scoring conditions are changed to award different
# points for a 3-0 win than a 5-0 win (or whatever), this code will break.

# compares against the defined validity conditions, returns false with the reason if any are not met.
def is_division_valid(prelim_record,prelim_count):
    school_count = prelim_record['School'].nunique()
    entry_count = len(prelim_record['Code'])
    if school_count < MINIMUM_SCHOOLS_PER_DIVISION:
        return [False,"Too few schools"]
    if entry_count < MINIMUM_TEAMS_PER_DIVISION:
        return [False,"Too few teams"]
    if prelim_count < MINIMUM_PRELIMS_PER_DIVISION:
        return [False,"Too few prelims"]
    return [True,"Division is valid"]

# Replaces elim rows with explicit walkovers, assigns a 3-0 win to the advancing entry.
def process_elim_walkovers(elim_record):
    elim_walkovers = pd.DataFrame()
    elim_walkovers = elim_record[elim_record['Win'].str.contains('advances')]
    if not elim_walkovers.empty:
        elim_walkovers['walkover_winner'] = elim_walkovers['Win'].str[:-9]
        elim_walkovers['aff_walks_over'] = elim_walkovers['walkover_winner'] == elim_walkovers['Aff']
        elim_walkovers['walkover_ballot'] = elim_walkovers['aff_walks_over'].replace({True: '3-0\tAFF', False: '3-0\tNEG'})
        elim_walkovers.drop(columns=['walkover_winner','aff_walks_over','Win'],axis=1,inplace=True)
        elim_walkovers['Win'] = elim_walkovers['walkover_ballot']
        elim_walkovers.drop('walkover_ballot',axis=1,inplace=True)
        elim_record.drop(elim_record.index[elim_record['Win'].str.contains('advances')],inplace=True)
        elim_record = pd.concat([elim_record,elim_walkovers[['Aff','Neg','Win']]])
    return elim_record

# Detects elim rounds where the tournament organizer has denoted an elim walkover by simply pairing two teams from the same
# school against each other and has not published a winner. Manual intervention is required to correct these elim rounds,
# because it's necessary to look at the next elim round to determine who actually advanced.
def detect_implied_walkovers(elim_record):
    elim_record['implied_walkover'] = (elim_record['Win'].isna()) & (~(elim_record['Neg'].isna())) & (~(elim_record['Aff'].isna()))
    return elim_record['implied_walkover'].any()

# replaces NaN and blanks with acceptable input for bye and walkover processing. 
# implicitly assumes that any team getting a bye is listed on the affirmative.
def eliminate_blank_teams(elim_record):
    elim_record
    na_values = {'Neg': 'BLANK_ENTRY', 'Win': 'Aff BYE'}
    elim_record.fillna(value=na_values,inplace=True)
    return elim_record

# awards a 0-3 loss for any team forfeiting an elimination round.
def process_elim_forfeits(elim_record):
    elim_forfeits = pd.DataFrame()
    elim_forfeits = elim_record[elim_record['Win'].str.contains('FFT')]
    if not elim_forfeits.empty:
        elim_forfeits[['aff_result','neg_result']] = elim_forfeits['Win'].str.split('\t+',expand=True)
        elim_forfeits['aff_result'] = elim_forfeits['aff_result'].str[4:]
        elim_forfeits['Win'] = elim_forfeits['aff_result'].replace({'FFT': '3-0\tNEG','BYE':'3-0\tAff'})
        elim_record = elim_record.drop(elim_record.index[elim_record['Win'].str.contains('FFT')])
        elim_record = pd.concat([elim_record,elim_forfeits])
    return elim_record

    
# awards a 3-0 to the team getting a bye
def process_elim_byes(elim_record):
    elim_byes = pd.DataFrame()
    elim_byes = elim_record[elim_record['Win'].str.contains('BYE')]
    if not elim_byes.empty:
        elim_byes[['Win','bye']] = elim_byes['Win'].str.split(' ',expand=True)
        elim_byes['Win'] = elim_byes['Win'].str.upper()
        elim_byes[['bye']] = elim_byes[['bye']].replace('BYE','3-0\t')
        elim_byes['Win'] = elim_byes['bye'] + elim_byes['Win']
        elim_byes = elim_byes.drop('bye',axis=1)
        elim_record = elim_record.drop(elim_record.index[elim_record['Win'].str.contains('BYE')])
        elim_record = pd.concat([elim_record,elim_byes])
    return elim_record

# ensure hybrid entires are not awarded points. There are currently no 'NDT-recognized' hybrid teams, if any exist this function 
# will need to be modified to account for this. That change would also force me to assign a school to the hybrid team.
def drop_hybrid_entries(tournament_points):
    tournament_points['is_team_hybrid'] = tournament_points['Code'].str.contains('/')
    tournament_points.drop(tournament_points[tournament_points.is_team_hybrid].index, inplace=True)
    tournament_points.drop(['is_team_hybrid'],axis=1,inplace=True)
    return tournament_points

def process_points_division(tournament_name,year,prelim_count,division):
    print_if_debug('\t'+division.name)
    school_division_points = pd.DataFrame()
    data_folder = 'tournament_results/'+str(year)+'/'+tournament_name
    prelimFilePath=data_folder+'/'+tournament_name+'-'+division.value+'-prelims.csv'
    tournament_prelims = pd.read_csv(prelimFilePath)
    [division_is_valid,validity_string] = is_division_valid(tournament_prelims,prelim_count)
    if division_is_valid:
        invalid_divisions_row=pd.DataFrame()
        tournament_prelims['prelim_winrate'] = tournament_prelims['Wins']/prelim_count
        tournament_prelims['prelim_points'] = tournament_prelims['prelim_winrate'].apply(ndt_points_from_prelims)
        tournament_points=tournament_prelims[['Code','School','prelim_points']]
        dir_list = os.listdir(data_folder)
        search_string = '-'+division.value+'-elim'
        elims_to_process = filter(lambda x: re.search(search_string, x), dir_list)
        elim_index=0
        for elim_filename in list(elims_to_process):
            print_if_debug('\t\t'+elim_filename)
            elim_index+=1
            points_column_header='elim_'+str(elim_index)+"_points"
            elim_record=pd.read_csv(data_folder+'/'+elim_filename)[['Aff','Neg','Win']]
            if detect_implied_walkovers(elim_record):
                raise ValueError("Human intervention needed: Implied walkover in ",elim_filename)
            elim_record = eliminate_blank_teams(elim_record)
            elim_record = process_elim_walkovers(elim_record)
            elim_record = process_elim_forfeits(elim_record)
            elim_record = process_elim_byes(elim_record)
            elim_record[['ballots','Win']] = elim_record['Win'].str.split('\t+',expand=True) # this is inelegant, but works
            elim_record[['winner_ballots','loser_ballots']] = elim_record['ballots'].str.split('-',expand=True) # breaks if there's not a dash in there, should be caught above
            elim_record[['loser_ballots']] = elim_record[['loser_ballots']].astype("int")
            elim_record[['winner_points']] = elim_record[['loser_ballots']].apply(ndt_winner_points_from_elims)
            elim_record[['loser_points']] = elim_record[['loser_ballots']].apply(ndt_loser_points_from_elims)
            elim_record['aff_win'] = elim_record['Win'].apply(lambda y: 1 if y=='AFF' else 0)## TODO: replace sad slow apply with vectorized happy fast 'replace'
            elim_record['neg_win'] = 1-elim_record['aff_win']
            elim_record['aff_points'] = elim_record['winner_points']*elim_record['aff_win']+elim_record['loser_points']*elim_record['neg_win']
            elim_record['neg_points'] = elim_record['winner_points']*elim_record['neg_win']+elim_record['loser_points']*elim_record['aff_win']
            temp_aff = pd.DataFrame()
            temp_neg = pd.DataFrame()
            temp_aff[['Code',points_column_header]] = elim_record[['Aff','aff_points']]
            temp_neg[['Code',points_column_header]] = elim_record[['Neg','neg_points']]
            elim_points = pd.concat([temp_aff,temp_neg])
            tournament_points = tournament_points.merge(elim_points,'left','Code')
            tournament_points[[points_column_header]] = tournament_points[[points_column_header]].fillna(0).astype(int)
        tournament_points = drop_hybrid_entries(tournament_points)
        tournament_points['total_points'] = tournament_points.drop(['Code','School'],axis=1).sum(axis=1)
        school_division_points = tournament_points[['School','total_points']].groupby('School',as_index=False).agg({'total_points': {lambda z: z.nlargest(MAXIMUM_ENTRIES_COUNTED_PER_SCHOOL).sum()}})
        school_division_points.columns = list(map(''.join, school_division_points.columns.values))
        school_division_points[tournament_name+'_'+division.value+'_points'] = school_division_points['total_points<lambda>']
        school_division_points = school_division_points.drop('total_points<lambda>',axis=1)
        apply_dictionary_to_results_dataframe(school_division_points,school_alias_dict)
    else:
        invalid_divisions_row=pd.DataFrame({'Tournament':[tournament_name],'Division':[division.name],'Reason':[validity_string]})
    return [school_division_points,invalid_divisions_row]
	
### Functions to split tournaments into divisions and integrate tournaments into one Big Table
def process_points_tournament(tournament,invalid_divisions_dataframe):
    tournament_name=tournament.name
    prelim_count_vector=tournament.prelim_counts
    division_vector=tournament.divisions
    year=tournament.year
    print_if_debug('Processing tournament: '+tournament_name)
    school_tournament_points=pd.DataFrame()
    for (division,prelim_count) in zip(division_vector, prelim_count_vector):
        if prelim_count==0:
            continue # don't process fake divisions
        division_points = pd.DataFrame()
        [division_points,division_invalid_row] = process_points_division(tournament_name,year,prelim_count,division)
        if school_tournament_points.empty:
            school_tournament_points = division_points # the merge will error out if there aren't any rounds in the division (and consequently the output dataframe is empty)
        elif not division_points.empty:
            school_tournament_points = school_tournament_points.merge(division_points,how='outer',on='School')
        if invalid_divisions_dataframe.empty:
            invalid_divisions_dataframe=division_invalid_row
        elif not division_invalid_row.empty:
            invalid_divisions_dataframe=pd.concat([invalid_divisions_dataframe,division_invalid_row],ignore_index=True)
    school_tournament_points.fillna(0,inplace=True)
    columns_to_add = school_tournament_points.loc[:,school_tournament_points.columns!='School'] # unsafe to reorder this list prior to merging
    if ~SEPARATE_ROUNDROBINS:
        non_varsity_rr_columns = [tournament_name+'_jv_points',tournament_name+'_n_points']
        varsity_rr_columns = [tournament_name+'_v_points',tournament_name+'_rr_points'] ## TODO: teach the Division class what it means to be varsity.
        varsity_plus_rr = columns_to_add.drop([x for x in non_varsity_rr_columns if x in columns_to_add.columns],axis=1)
        new_varsity_points=varsity_plus_rr.sum(axis=1)
        if ~new_varsity_points.empty:
            columns_to_add = columns_to_add.drop([x for x in varsity_rr_columns if x in columns_to_add.columns],axis=1)
            columns_to_add[tournament_name+'_v_points'] = new_varsity_points
    total_tournament_points = columns_to_add.sum(axis=1)
    school_tournament_points[tournament_name+'_total_points'] = total_tournament_points
    return [school_tournament_points,invalid_divisions_dataframe]

def tournament_merge(cumulative_list,new_tournament):
    if cumulative_list.empty:
        return new_tournament
    elif new_tournament.empty:
        return cumulative_list # the merge will error out if this is the first tourney processed because there's no 'school' column
    new_cumulative_list = cumulative_list.merge(new_tournament,how='outer',on='School')
    new_cumulative_list.fillna(0,inplace=True)
    return new_cumulative_list

def sum_legal_tournaments(cumulative_points,division,legal_tournament_count):
    column_label_substring = '_'+division.value+'_'
    filtered_cumulative_points = cumulative_points.filter(like=column_label_substring)
    filtered_cumulative_points = filtered_cumulative_points.apply(pd.Series.nlargest,axis=1,n=legal_tournament_count) # this is slow, but i don't know a faster way
    filtered_cumulative_points.fillna(0,inplace=True)
    total_points=pd.DataFrame()
    total_points['School'] = cumulative_points['School']
    total_points[division.value+'_total_points'] = filtered_cumulative_points.sum(axis=1)
    return total_points


### validate that required folders exist
    print_if_debug('looking for '+REPORTS_GENERATED_FOLDER+'...')
if not os.path.exists(REPORTS_GENERATED_FOLDER):
    print_if_debug('no reports folder -- creating it...')
    os.makedirs(REPORTS_GENERATED_FOLDER)
    if not os.path.exists(REPORTS_GENERATED_FOLDER):
        raise Exception("Unable to create directory ."+REPORTS_GENERATED_FOLDER)


### define tournaments and execute
tournament_list=pd.read_csv('tournaments-'+str(YEAR_TO_PROCESS)+'.csv')
if REPORT_TO_GENERATE==1:
    tournament_list=tournament_list[tournament_list['season']=='fall']

invalid_divisions_list=pd.DataFrame()
cumulative_points = pd.DataFrame()
for tournament_index,tournament_data in tournament_list.iterrows():
    division_rounds = [tournament_data['varsity_rounds'],tournament_data['junior_varsity_rounds'],tournament_data['novice_rounds'],tournament_data['round_robin_rounds']]
    divisions = [Division.VARSITY,Division.JUNIOR_VARSITY,Division.NOVICE,Division.ROUND_ROBIN]
    tournament_to_process=tournament(tournament_data['tournament_name'],YEAR_TO_PROCESS,division_rounds,divisions)
    [tournament_points,invalid_divisions_list]=process_points_tournament(tournament_to_process,invalid_divisions_list)
    cumulative_points = tournament_merge(cumulative_points,tournament_points)
if not invalid_divisions_list.empty:
    print_if_debug('invalid divisions:')
    print_if_debug(invalid_divisions_list.to_string())
    invalid_divisions_list.to_csv(index=False,path_or_buf=REPORTS_GENERATED_FOLDER+"invalid_divisions_"+str(YEAR_TO_PROCESS)+"_"+report_season+".csv")
### report generation
def add_rank_column(dataframe):
    dataframe['Rank'] = range(1,len(dataframe)+1) # can't just return the index, it starts at zero.
    dataframe['Rank'] = dataframe['Rank'].astype(str)+'.'
    columns = dataframe.columns.tolist()
    columns = columns[-1:] + columns[:-1]
    return dataframe[columns]

total_points_column = sum_legal_tournaments(cumulative_points,Division.TOTAL,MAXIMUM_TOURNAMENTS_COUNTED_PER_SCHOOL)
varsity_points_column = sum_legal_tournaments(cumulative_points,Division.VARSITY,MAXIMUM_TOURNAMENTS_COUNTED_PER_SCHOOL)
sweepstakes_results_for_reports = pd.DataFrame()
sweepstakes_results_for_reports = total_points_column.merge(varsity_points_column,how='outer',on='School').fillna(0) # some schools don't run non-varsity teams, some only run non-varsity teams
sweepstakes_results_for_reports['NDT pts'] = sweepstakes_results_for_reports.total_total_points.astype(int) # decimal points are big ugly
sweepstakes_results_for_reports['Varsity pts'] = sweepstakes_results_for_reports.v_total_points.astype(int)
sweepstakes_results_for_reports.drop(columns=['v_total_points','total_total_points'],inplace=True) # gotta rename the column, gotta remove decimal points, may as well permute.

# sweepstakes_results_for_reports = apply_dictionary_to_results_dataframe(sweepstakes_results_for_reports,school_alias_dict) #must replace school names prior to matching by the display name

schools_by_districts = pd.read_csv('ndt-districts.csv')
community_colleges = pd.read_csv('community-colleges.csv')
sweepstakes_results_for_reports = sweepstakes_results_for_reports.merge(schools_by_districts,how='left',on='School').fillna(-1) # we should put in a placeholder district for undistricted schools and then give up
undistricted_schools=sweepstakes_results_for_reports[sweepstakes_results_for_reports['District'] == -1]
if not undistricted_schools.empty:
    print('The following schools are not assigned to a district:')
    print(undistricted_schools.to_string())
    raise Exception("There are schools without a district.")
sweepstakes_results_for_reports = sweepstakes_results_for_reports.merge(community_colleges,how='left',on='School') # but we should not assume every school is listed as a non-CC in the community-colleges file.
sweepstakes_results_for_reports.fillna(value=False,inplace=True)
sweepstakes_results_for_reports['CC'].replace({True: 'Y', False: 'N'},inplace=True) # i want to display this in a pretty way.



sweepstakes_results_for_reports.to_csv(index=False,path_or_buf=REPORTS_GENERATED_FOLDER+"sweepstakes_output_"+str(YEAR_TO_PROCESS)+"_"+report_season+"_full.csv")


## Remove non-NDT-members from spring report tabulation, but still record them.

if REPORT_TO_GENERATE==2:
    ndt_members = pd.read_csv('ndt-members.csv')
    ndt_members_current = pd.DataFrame()
    ndt_members_current[['School','Member']] = ndt_members[['Display_School',str(YEAR_TO_PROCESS)]]
    
    sweepstakes_results_for_reports = sweepstakes_results_for_reports.merge(ndt_members_current,how='left',on='School').fillna(-1)
    schools_without_membership_info = sweepstakes_results_for_reports[sweepstakes_results_for_reports['Member']==-1]
    if not schools_without_membership_info.empty:
        schools_without_membership_info = schools_without_membership_info['School']
        print('the following schools lack membership status for '+str(YEAR_TO_PROCESS)+':')
        print(schools_without_membership_info.to_string())
        raise Exception('There are schools lacking membership status')
    sweepstakes_results_for_reports = sweepstakes_results_for_reports[sweepstakes_results_for_reports['Member']==1]
    sweepstakes_results_for_reports.drop(columns=['Member'],axis=1,inplace=True)



sweepstakes_top10_overall = add_rank_column(sweepstakes_results_for_reports.sort_values('NDT pts',ascending=False,ignore_index=True).head(10))
sweepstakes_top10_varsity = add_rank_column(sweepstakes_results_for_reports.sort_values('Varsity pts',ascending=False,ignore_index=True).head(10))
sweepstakes_top10_overall_CC = add_rank_column(sweepstakes_results_for_reports[sweepstakes_results_for_reports['CC']=='Y'].sort_values('NDT pts',ascending=False,ignore_index=True))
sweepstakes_overall_rankings = add_rank_column(sweepstakes_results_for_reports.sort_values('NDT pts',ascending=False,ignore_index=True))
sweepstakes_varsity_rankings = add_rank_column(sweepstakes_results_for_reports.sort_values('Varsity pts',ascending=False,ignore_index=True))

district_overall_sweepstakes_points = {}
for district in NDT_DISTRICTS: # filter by district, then sort by 'overall', then add a rank column, then it's good
    district_overall_sweepstakes_points[district] = add_rank_column(sweepstakes_results_for_reports[sweepstakes_results_for_reports['District']==district].sort_values('NDT pts',ascending=False,ignore_index=True))

if (REPORT_TO_GENERATE==2) & (~NO_REPORT_GEN):
    if not os.path.exists(PREVIOUS_YEAR_REPORTS_FOLDER):
        raise Exception("Could not find previous year reports -- Do they exist in "+PREVIOUS_YEAR_REPORTS_FOLDER+"?")
    last_fall_filename = PREVIOUS_YEAR_REPORTS_FOLDER+'sweepstakes_output_'+str(YEAR_TO_PROCESS-1)+'_fall_full.csv'
    last_fall_results = pd.read_csv(last_fall_filename)
    last_spring_filename = PREVIOUS_YEAR_REPORTS_FOLDER+'sweepstakes_output_'+str(YEAR_TO_PROCESS-1)+'_spring_full.csv'
    last_spring_results = pd.read_csv(last_spring_filename)
    sweepstakes_results_for_reports['new-schools-eligible'] = sweepstakes_results_for_reports['NDT pts']>=NEW_SCHOOL_POINTS_THRESHOLD
    sweepstakes_results_for_reports['existed_last_fall'] = sweepstakes_results_for_reports['School'].isin(last_fall_results['School'])
    sweepstakes_results_for_reports['existed_last_spring'] = sweepstakes_results_for_reports['School'].isin(last_spring_results['School'])
    if ~(sweepstakes_results_for_reports['existed_last_fall'].all()): # don't process it if there aren't any new schools
        new_schools_maximally_permissive = sweepstakes_results_for_reports[~((sweepstakes_results_for_reports['existed_last_fall']) & (sweepstakes_results_for_reports['existed_last_spring']))]
        new_schools_maximally_permissive.drop(columns=['new-schools-eligible'],axis=1,inplace=True)
        new_schools_for_reports=sweepstakes_results_for_reports[(sweepstakes_results_for_reports['new-schools-eligible']) & (~sweepstakes_results_for_reports['existed_last_fall'])]
        new_schools_for_reports.drop(columns=['new-schools-eligible','existed_last_fall','existed_last_spring'],axis=1,inplace=True)
        new_schools_for_reports = add_rank_column(new_schools_for_reports.sort_values('NDT pts',ascending=False,ignore_index=True))
        print_if_debug("new schools:")
        print_if_debug(new_schools_for_reports.to_string())
    else:
        new_schools_for_reports=pd.DataFrame()
    if VALIDATION_MODE:
        last_spring_results['exists_next_spring'] = last_spring_results['School'].isin(sweepstakes_results_for_reports['School'])
        vanishing_schools = last_spring_results[~last_spring_results['exists_next_spring']]
        vanishing_schools = vanishing_schools.reindex(columns=['School'])
        print('Vanishing schools:')
        print(vanishing_schools.to_string())
        print('New schools:')
        print(new_schools_maximally_permissive['School'].to_string())
        print('If any schools appear on both lists with different names, update the aliases file.')
        last_spring_results.drop(columns=['exists_next_spring'],axis=1,inplace=True)
    
    movers_for_reports=sweepstakes_results_for_reports[sweepstakes_results_for_reports['existed_last_spring']] ## don't want to try and calculate NA last-year pts
    last_year_just_points=last_spring_results[['School','NDT pts']]
    movers_for_reports=movers_for_reports.merge(last_year_just_points,how='left',on='School',suffixes=('_current','_previous'))
    movers_for_reports.drop(columns=['Varsity pts','District','CC','new-schools-eligible','existed_last_spring'],axis=1,inplace=True)
    movers_for_reports['Moved']=movers_for_reports['NDT pts_current']-movers_for_reports['NDT pts_previous']
    movers_for_reports['NDT pts'] = movers_for_reports['NDT pts_current']
    movers_for_reports.drop(columns=['NDT pts_current','NDT pts_previous'],axis=1,inplace=True)
    movers_for_reports=movers_for_reports.reindex(columns=['School','NDT pts','Moved'])
    movers_for_reports = movers_for_reports[movers_for_reports['Moved']>=MOVERS_THRESHOLD]
    if ~movers_for_reports.empty:
        movers_for_reports=add_rank_column(movers_for_reports.sort_values('Moved',ascending=False,ignore_index=True))
    print_if_debug("movers:")
    print_if_debug(movers_for_reports.to_string())






## output to word tables



season_caps=report_season.upper()
season_sentence=report_season.capitalize()
report_replacement_dictionary={"$YEAR":report_year,"$FIRST":report_ordinal,"$SEASON_LOWER":report_season,"$SEASON_UPPER":season_caps,"$SEASON_SENTENCE":season_sentence}

def report_update_year(template_document):
    for paragraph in template_document.paragraphs:
        for replacement_item in report_replacement_dictionary:
            if paragraph.text.find(replacement_item)>=0:
                runs=paragraph.runs
                for i in range(len(runs)):
                    if runs[i].text.find(replacement_item)>=0:
                        runs[i].text = runs[i].text.replace(replacement_item,report_replacement_dictionary[replacement_item])
    for replacement_item in report_replacement_dictionary:
        runs=template_document.sections[0].footer.paragraphs[0].runs
        for i in range(len(runs)):
            if runs[i].text.find(replacement_item)>=0:
                runs[i].text = runs[i].text.replace(replacement_item,report_replacement_dictionary[replacement_item])
        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.find(replacement_item)>=0:
                            runs=paragraph.runs
                            for i in range(len(runs)):
                                if runs[i].text.find(replacement_item)>=0:
                                    runs[i].text = runs[i].text.replace(replacement_item,report_replacement_dictionary[replacement_item])
    return template_document
        

def append_word_results_table(results_document,results_dataframe,append_footers):
    results_column_widths=[0.5,3,0.7,0.9,0.7,0.4]
    created_table = results_document.add_table(results_dataframe.shape[0]+1,results_dataframe.shape[1],style="NDTSweepstakes")

    for j in range(results_dataframe.shape[-1]):
        created_table.cell(0,j).text = results_dataframe.columns[j]
        created_table.cell(0,j).width=docx.shared.Inches(results_column_widths[j]) # this is disgusting, but word doesn't respect widths set to entire columns.
    
    for i in range(results_dataframe.shape[0]):
        for j in range(results_dataframe.shape[-1]):
            created_table.cell(i+1,j).text=str(results_dataframe.values[i,j])
            created_table.cell(i+1,j).width=docx.shared.Inches(results_column_widths[j]) # i hate it too.
    

    if append_footers:
        results_document.add_paragraph('')
        footer_table=results_document.add_table(1,1,style="NDTSweepstakes") # in the default format, that's just a blue block across the bottom of the table.
        results_document.add_paragraph('')
    return results_document

def append_table_header(results_document,title_string):
    results_document.add_heading(title_string,level=3)
    results_document.add_paragraph('')
    return results_document
    
if NO_REPORT_GEN:
    print_if_debug('no reports generated...')
else:
    print_if_debug('creating word tables...')
    results_document = docx.Document('sweepstakes-table-template.docx')
    results_document = report_update_year(results_document)
    
    print_if_debug('updating top-10...')
    results_document = append_table_header(results_document,"Top 10 Overall Rankings")
    results_document = append_word_results_table(results_document,sweepstakes_top10_overall,True)
    
    results_document = append_table_header(results_document,"Top 10 Varsity Rankings")
    results_document = append_word_results_table(results_document,sweepstakes_top10_varsity,True)
    
    print_if_debug('updating CCs...')
    results_document = append_table_header(results_document,"Top CC Rankings")
    results_document = append_word_results_table(results_document,sweepstakes_top10_overall_CC,True)
    results_document.add_page_break()
    
    if REPORT_TO_GENERATE==2:
        print_if_debug('updating new schools...')
        results_document = append_table_header(results_document,"New Schools")
        results_document.add_paragraph('New schools with '+str(NEW_SCHOOL_POINTS_THRESHOLD)+' or more Overall NDT points (new schools are those schools that did not earn points fall of the previous years):')
        if new_schools_for_reports.empty:
            results_document.add_paragraph('\tAccording to our records, there were no new schools that were '+str(YEAR_TO_PROCESS)+'-'+str((YEAR_TO_PROCESS+1)%100)+' NDT subscribers.').bold = True
        else:
            results_document = append_word_results_table(results_document,new_schools_for_reports,True)
        print_if_debug('updating movers...')
        results_document = append_table_header(results_document,"Movers")
        results_document.add_paragraph('Movers with '+str(MOVERS_THRESHOLD)+' or more Overall NDT points than the previous year (comparing the Spring reports; schools who were not members the previous year are not eligible):')
        if movers_for_reports.empty:
            results_document.add_paragraph('\tAccording to our records, there were no schools that moved by '+str(MOVERS_THRESHOLD)+'Overall NDT points.').bold = True
        else:
            results_document = append_word_results_table(results_document,movers_for_reports,True)
        
    
    
    print_if_debug('updating full overall...')
    results_document = append_table_header(results_document,"Overall Rankings")
    results_document = append_word_results_table(results_document,sweepstakes_overall_rankings,True)
    results_document.add_page_break()
    
    print_if_debug('updating full varsity...')
    results_document = append_table_header(results_document,"Varsity Rankings")
    results_document = append_word_results_table(results_document,sweepstakes_varsity_rankings,True)
    results_document.add_page_break()
    
    print_if_debug('printing division tables...')
    results_document = append_table_header(results_document,"Overall Rankings by District")
    for district in NDT_DISTRICTS:
        results_document = append_word_results_table(results_document,district_overall_sweepstakes_points[district],False)
        results_document.add_paragraph('')
    footer_table=results_document.add_table(1,1,style="NDTSweepstakes")
    results_document.add_paragraph('')
    results_document.add_page_break()
    
    print_if_debug('adding appendices...')
    results_composer=Composer(results_document)
    procedure_document=docx.Document('sweepstakes-procedure.docx')
    
    results_composer.append(procedure_document)
    
    print_if_debug('saving...')
    report_filename=REPORTS_GENERATED_FOLDER+str(YEAR_TO_PROCESS)+'-'+str((YEAR_TO_PROCESS+1)%100)+'-NDT-Points-Standings-'+season_sentence+'.docx'
    results_composer.save(report_filename)
print_if_debug('done!')